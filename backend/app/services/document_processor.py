"""
Document processor — loads uploaded files and splits them into chunks.

Supported formats: PDF, TXT, Markdown, DOCX.
Uses LangChain's RecursiveCharacterTextSplitter for intelligent chunking.
"""

from __future__ import annotations

import logging
from pathlib import Path

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredMarkdownLoader,
)
from langchain_core.documents import Document

logger = logging.getLogger(__name__)

# ── Supported file extensions ──────────────────────────────────────────────────
SUPPORTED_EXTENSIONS: set[str] = {".pdf", ".txt", ".md", ".docx"}

# ── Chunking defaults ─────────────────────────────────────────────────────────
DEFAULT_CHUNK_SIZE = 1000
DEFAULT_CHUNK_OVERLAP = 200


class DocumentProcessor:
    """Load files and split them into semantically meaningful chunks."""

    def __init__(
        self,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
    ) -> None:
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

    # ── Public API ─────────────────────────────────────────────────────────────

    def process_file(self, file_path: Path) -> list[Document]:
        """
        Load a single file and split it into chunks.

        Args:
            file_path: Absolute path to the uploaded file.

        Returns:
            List of LangChain Document objects with metadata.

        Raises:
            ValueError: If the file extension is not supported.
            FileNotFoundError: If the file does not exist.
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = file_path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: '{ext}'. "
                f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        logger.info("Processing file: %s (type: %s)", file_path.name, ext)

        # Load raw documents
        raw_docs = self._load_file(file_path, ext)

        # Add source metadata
        for doc in raw_docs:
            doc.metadata["source"] = file_path.name

        # Split into chunks
        chunks = self.text_splitter.split_documents(raw_docs)
        logger.info(
            "File '%s' produced %d chunks from %d pages/sections",
            file_path.name,
            len(chunks),
            len(raw_docs),
        )
        return chunks

    # ── Private helpers ────────────────────────────────────────────────────────

    def _load_file(self, file_path: Path, ext: str) -> list[Document]:
        """Dispatch to the correct loader based on file extension."""
        if ext == ".pdf":
            return PyPDFLoader(str(file_path)).load()

        if ext == ".txt":
            return TextLoader(str(file_path), encoding="utf-8").load()

        if ext == ".md":
            return UnstructuredMarkdownLoader(str(file_path)).load()

        if ext == ".docx":
            return self._load_docx(file_path)

        # Should never reach here due to the extension check above
        raise ValueError(f"No loader available for extension: {ext}")

    @staticmethod
    def _load_docx(file_path: Path) -> list[Document]:
        """Load a DOCX file using python-docx."""
        from docx import Document as DocxDocument  # type: ignore[import-untyped]

        doc = DocxDocument(str(file_path))
        full_text = "\n".join(
            paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()
        )
        return [
            Document(
                page_content=full_text,
                metadata={"source": file_path.name},
            )
        ]


# ── Module-level singleton ─────────────────────────────────────────────────────
document_processor = DocumentProcessor()
